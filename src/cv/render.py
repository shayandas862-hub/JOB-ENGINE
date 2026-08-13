"""Render a tailored CV to a single-column, ATS-safe .docx.

ATS parsers choke on tables, text boxes, columns, and images — so this uses none
of them: a single column of real paragraphs, real headings, and real bullet
lists (never a literal '•'). Contact details sit in the body, not a header.
Input is the gated PhrasedBlocks (truth-checked) grouped into fixed sections;
output is deterministic bytes, so the golden test stays stable.
"""
from __future__ import annotations

import io
from dataclasses import dataclass

from docx import Document

# CV sections in order; each maps to the cv_block kinds that feed it.
SECTION_ORDER = (
    ("Experience", {"role", "achievement"}),
    ("Skills", {"skill_evidence"}),
    ("Education", {"education"}),
)


@dataclass(frozen=True)
class CvHeader:
    name: str
    contact: str = ""          # one line, e.g. 'email | phone | London'


def _subhead(block) -> str:
    """The bold line above a block's bullet: 'Title — Organisation (dates)'."""
    left = " — ".join(part for part in (block.title, block.organisation) if part)
    if block.date_range:
        return f"{left} ({block.date_range})" if left else f"({block.date_range})"
    return left


def build_document(header: CvHeader, blocks):
    """Assemble the python-docx Document (single column, no tables)."""
    doc = Document()
    doc.add_paragraph(header.name, style="Title")
    if header.contact:
        doc.add_paragraph(header.contact)

    for section_title, kinds in SECTION_ORDER:
        section_blocks = [b for b in blocks if b.kind in kinds]
        if not section_blocks:
            continue
        doc.add_heading(section_title, level=1)
        for block in section_blocks:
            subhead = _subhead(block)
            if subhead:
                doc.add_paragraph().add_run(subhead).bold = True
            doc.add_paragraph(block.bullet, style="List Bullet")
    return doc


def render_cv(header: CvHeader, blocks) -> bytes:
    """Render the CV to .docx bytes (ready to attach to a Notion card)."""
    buf = io.BytesIO()
    build_document(header, blocks).save(buf)
    return buf.getvalue()
