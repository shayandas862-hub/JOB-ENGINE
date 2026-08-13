"""Tests for src/cv/render — single-column, ATS-safe .docx.

Golden test: render a known CV, read it back, and pin the exact (style, text)
paragraph sequence. Plus ATS-safety guards (no tables, one column) and byte
determinism (same input -> same bytes), which is what makes the golden stable.
"""
from __future__ import annotations

import io

from docx import Document


def pb(block_id, kind, title, organisation, date_range, bullet):
    from cv.phrase import PhrasedBlock
    return PhrasedBlock(block_id=block_id, kind=kind, title=title, organisation=organisation,
                        date_range=date_range, source_fact=bullet, bullet=bullet)


def _blocks():
    return [
        pb(1, "role", "Senior Data Analyst", "Acme", "2021-2023",
           "Led analytics, cutting reporting time 40%."),
        pb(2, "achievement", None, None, None, "Rebuilt the KPI pipeline end to end."),
        pb(3, "skill_evidence", None, None, None, "SQL, Python and dbt across production pipelines."),
        pb(4, "education", "BSc Computer Science", "University of X", "2016-2019",
           "Graduated with first-class honours."),
    ]


def _header():
    from cv.render import CvHeader
    return CvHeader("Shayan Das", "shayan@example.com | London")


def _read(data):
    return Document(io.BytesIO(data))


def test_render_produces_the_golden_ats_structure():
    from cv.render import render_cv
    doc = _read(render_cv(_header(), _blocks()))
    seq = [(p.style.name, p.text) for p in doc.paragraphs if p.text.strip()]
    assert seq == [
        ("Title", "Shayan Das"),
        ("Normal", "shayan@example.com | London"),
        ("Heading 1", "Experience"),
        ("Normal", "Senior Data Analyst — Acme (2021-2023)"),
        ("List Bullet", "Led analytics, cutting reporting time 40%."),
        ("List Bullet", "Rebuilt the KPI pipeline end to end."),
        ("Heading 1", "Skills"),
        ("List Bullet", "SQL, Python and dbt across production pipelines."),
        ("Heading 1", "Education"),
        ("Normal", "BSc Computer Science — University of X (2016-2019)"),
        ("List Bullet", "Graduated with first-class honours."),
    ]


def test_render_is_ats_safe_single_column_no_tables():
    from docx.oxml.ns import qn

    from cv.render import render_cv
    doc = _read(render_cv(_header(), _blocks()))
    assert doc.tables == []                               # no tables — ATS-safe
    assert len(doc.sections) == 1                         # single section...
    cols = doc.sections[0]._sectPr.find(qn("w:cols"))    # ...single column
    assert cols is None or cols.get(qn("w:num")) in (None, "1")


def test_render_is_deterministic():
    from cv.render import render_cv
    a, b = render_cv(_header(), _blocks()), render_cv(_header(), _blocks())
    assert isinstance(a, bytes) and len(a) > 0
    # .docx bytes carry a zip timestamp, so compare the read-back structure.
    seq = lambda data: [(p.style.name, p.text) for p in _read(data).paragraphs]
    assert seq(a) == seq(b)                              # same input -> same document


def test_empty_sections_are_omitted():
    from cv.render import render_cv
    only_roles = [pb(1, "role", "Analyst", "Acme", "2021", "Did analytics.")]
    doc = _read(render_cv(_header(), only_roles))
    headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
    assert headings == ["Experience"]                    # no Skills/Education headings


def test_block_without_a_subhead_renders_only_its_bullet():
    from cv.render import render_cv
    doc = _read(render_cv(_header(), [pb(1, "skill_evidence", None, None, None, "SQL and Python.")]))
    styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
    # Title, contact, Skills heading, then just the bullet — no bold sub-line
    assert styles == ["Title", "Normal", "Heading 1", "List Bullet"]
